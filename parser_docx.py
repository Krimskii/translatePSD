from docx import Document


def iter_docx_text_containers(doc):
    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def parse_docx(file):
    doc = Document(file)
    return [container.text for container in iter_docx_text_containers(doc)]
